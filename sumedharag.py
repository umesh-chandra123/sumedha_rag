import os
from typing import List, Tuple
import streamlit as st
import io

from dotenv import load_dotenv
from huggingface_hub import snapshot_download
from pypdf import PdfReader 
from docx import Document as DocxDocument

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough


# ---------------- boot ----------------
load_dotenv()
st.set_page_config(page_title="TXT → RAG (Gemini + FAISS + HF)", page_icon="🧠", layout="wide")
st.title("TXT → RAG (Gemini + FAISS + HF)")

# constants
HF_MODEL_ID = "sentence-transformers/all-MiniLM-L6-v2"
DEFAULT_LOCAL_DIR = os.getenv("LOCAL_EMB_DIR", "./models/all-MiniLM-L6-v2")
FIXED_TEMPERATURE = 0.1


# ---------------- sidebar ----------------
with st.sidebar:
    st.subheader("Keys & Models")

    # Gemini key: stored to env var Gemini wrappers expect
    gemini_key = st.text_input(
        "GOOGLE_API_KEY (or GEMINI_API_KEY)",
        type="password",
        value=os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY") or "",
    )
    if gemini_key:
        os.environ["GOOGLE_API_KEY"] = gemini_key

    # HF token: helps with rate limits / gated repos for the FIRST download
    hf_token = st.text_input(
        "Hugging Face token (optional—for first download)",
        type="password",
        value=os.getenv("HUGGINGFACE_HUB_TOKEN") or "",
    )
    if hf_token:
        os.environ["HUGGINGFACE_HUB_TOKEN"] = hf_token

    gemini_model = st.selectbox(
        "Gemini model",
        ["gemini-2.5-flash", "gemini-2.5-pro", "gemini-flash-latest", "gemini-1.5-flash"],
        index=0,
    )
    st.caption(f"Temperature: *{FIXED_TEMPERATURE}*")

    st.divider()
    st.subheader("Embeddings")
    st.caption("First run will download the model to this folder; later runs load from disk.")
    local_dir = st.text_input("Local embedding folder", value=DEFAULT_LOCAL_DIR)
    laptop_mode = st.toggle("Laptop-friendly mode", value=True)

    st.divider()
    st.subheader("Retrieval")
    k = st.slider("Top-K chunks", 1, 10, 4)
    mmr = st.toggle("Use MMR (diverse results)", value=True)

    st.divider()
    st.subheader("Chunking")
    chunk_mode = st.radio("Mode", ["Auto (recommended)", "Manual"], index=0)
    if chunk_mode == "Manual":
        chunk_size = st.number_input("Chunk size (chars)", 128, 4000, 800, step=64)
        chunk_overlap = st.number_input("Chunk overlap (chars)", 0, 1000, 120, step=20)
    else:
        chunk_size = None
        chunk_overlap = None

    st.divider()
    if st.button("Clear FAISS index", type="secondary"):
        st.session_state.pop("vs", None)
        st.session_state.pop("retriever", None)
        st.success("Index cleared.")


# ---------------- state ----------------
if "messages" not in st.session_state:
    st.session_state.messages = []
if "vs" not in st.session_state:
    st.session_state.vs = None
if "retriever" not in st.session_state:
    st.session_state.retriever = None


# ---------------- helpers ----------------
def have_local_model(folder: str) -> bool:
    if not os.path.isdir(folder):
        return False
    needed = ["config.json", "tokenizer.json"]
    present = all(os.path.exists(os.path.join(folder, f)) for f in needed)
    has_weights = any(
        os.path.exists(os.path.join(folder, f))
        for f in ("model.safetensors", "pytorch_model.bin")
    )
    return present and has_weights


@st.cache_resource(show_spinner=False)
def ensure_local_model(model_id: str, folder: str, token: str | None) -> str:
    """
    If the folder already has the model, return it.
    Otherwise attempt a one-time download to that folder.
    """
    if have_local_model(folder):
        return folder

    os.makedirs(folder, exist_ok=True)

    # Make sure we don't end up with symlinks that break container copies
    snapshot_download(
        repo_id=model_id,
        local_dir=folder,
        local_dir_use_symlinks=False,
        token=token or os.getenv("HUGGINGFACE_HUB_TOKEN"),
    )
    return folder


def suggest_chunk_params(laptop: bool) -> Tuple[int, int]:
    # tuned to all-MiniLM-L6-v2 (384d) for general text; keep it simple + robust
    size, overlap = (900, 120)
    if laptop:
        size = 700
        overlap = 100
    return size, overlap


def build_splitter() -> RecursiveCharacterTextSplitter:
    if chunk_mode == "Manual":
        return RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    size, overlap = suggest_chunk_params(laptop_mode)
    return RecursiveCharacterTextSplitter(chunk_size=size, chunk_overlap=overlap)


def docs_from_files(files) -> List[Document]:
    docs = []
    for f in files:
        file_name = f.name
        raw_text = ""
        
        if file_name.lower().endswith(".txt"):
            raw_text = f.read().decode("utf-8", errors="ignore")
            
        elif file_name.lower().endswith(".pdf"):
            try:
                f.seek(0)
                pdf_reader = PdfReader(io.BytesIO(f.read()))
                text_parts = [page.extract_text() for page in pdf_reader.pages if page.extract_text()]
                raw_text = "\n\n".join(text_parts)
            except Exception as e:
                st.error(f"Error reading PDF {file_name}: Check if the PDF is text-searchable. Error: {e}")
                continue

        elif file_name.lower().endswith(".docx"):
            try:
                f.seek(0)
                doc = DocxDocument(io.BytesIO(f.read()))
                raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except Exception as e:
                st.error(f"Error reading DOCX {file_name}: {e}")
                continue
            
        else:
            st.warning(f"Skipping unsupported file type: {file_name}")
            continue
            
        if raw_text:
            docs.append(Document(page_content=raw_text, metadata={"source": file_name}))
            
    return docs


def format_docs(docs: List[Document]) -> str:
    return "\n\n".join(f"[{d.metadata.get('source','')}]\n{d.page_content}" for d in docs)


def ensure_retriever():
    if st.session_state.retriever is not None:
        return st.session_state.retriever
    if st.session_state.vs is None:
        return None
    if mmr:
        fetch_k = max(k * (2 if laptop_mode else 4), 10)
        retr = st.session_state.vs.as_retriever(search_type="mmr", search_kwargs={"k": k, "fetch_k": fetch_k})
    else:
        retr = st.session_state.vs.as_retriever(search_kwargs={"k": k})
    st.session_state.retriever = retr
    return retr


@st.cache_resource(show_spinner=False)
def load_embeddings_from_folder(folder: str, normalize: bool = True) -> HuggingFaceEmbeddings:
    # lock into offline once files exist
    os.environ["HF_HUB_OFFLINE"] = "1"
    os.environ["TRANSFORMERS_OFFLINE"] = "1"
    return HuggingFaceEmbeddings(
        model_name=folder,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": normalize},
    )


# ---------------- ingest ----------------
st.subheader("Upload .txt files")
files = st.file_uploader(
    "Upload text, PDF, or Word files",
    type=["txt", "pdf", "docx"],
    accept_multiple_files=True
)


c1, c2 = st.columns([1, 1])
with c1:
    if st.button("Build / Rebuild Index", type="primary", disabled=(not files)):
        # 1) ensure local embedding model (download on first run only)
        try:
            with st.spinner("Preparing embedding model (first time may download)…"):
                folder = ensure_local_model(HF_MODEL_ID, local_dir, hf_token or None)
        except Exception as e:
            st.error(
                "Could not download the embedding model. "
                "If you're behind a corporate proxy/SSL interceptor, "
                "download on a machine with internet and copy the folder to this path:\n"
                f"{os.path.abspath(local_dir)}\n\n"
                f"Error: {e}"
            )
            st.stop()

        # 2) load embeddings strictly from disk
        embeddings = load_embeddings_from_folder(folder)

        # 3) split and index
        splitter = build_splitter()
        raw_docs = docs_from_files(files)
        chunks = splitter.split_documents(raw_docs)
        with st.spinner("Embedding & indexing…"):
            st.session_state.vs = FAISS.from_documents(chunks, embeddings)
            st.session_state.retriever = None
        st.success(f"Indexed {len(chunks)} chunks from {len(files)} file(s).")

with c2:
    if have_local_model(local_dir):
        st.info(f"Local embedding ready ✅  ({HF_MODEL_ID})")
    else:
        st.warning("Local embedding not present yet. It will be downloaded on first index build.")


# ---------------- LLM & chain ----------------
if not (os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")):
    st.warning("Add your *GOOGLE_API_KEY* (or *GEMINI_API_KEY*) in the sidebar to query Gemini.")
else:
    # model init with a small fallback
    try:
        llm = ChatGoogleGenerativeAI(model=gemini_model, temperature=FIXED_TEMPERATURE)
    except Exception:
        st.toast("Selected model not available here, using gemini-1.5-flash.", icon="⚠️")
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=FIXED_TEMPERATURE)

    prompt = ChatPromptTemplate.from_messages(
        [
            ("system",
             "Use the provided context to answer. If the answer isn't in the context, say you don't know. "
             "Include short file citations like [filename.txt] when you rely on a chunk."),
            ("human", "Question:\n{question}\n\nContext:\n{context}\n\nAnswer:")
        ]
    )

    retriever = ensure_retriever()

    def make_chain():
        if retriever is None:
            return None
        return (
            {
                "context": retriever | (lambda ds: format_docs(ds)),
                "question": RunnablePassthrough(),
            }
            | prompt
            | llm
            | StrOutputParser()
        )

    chain = make_chain()

    # ---------------- chat ----------------
    st.subheader("Chat over your files")

    for role, content in st.session_state.messages:
        with st.chat_message(role):
            st.markdown(content)

    user_q = st.chat_input("Ask a question about your uploaded .txt files…")
    if user_q:
        with st.chat_message("user"):
            st.markdown(user_q)
        st.session_state.messages.append(("user", user_q))

        if chain is None:
            with st.chat_message("assistant"):
                st.warning("Please build the index first.")
        else:
            with st.chat_message("assistant"):
                holder = st.empty()
                final = []
                for chunk in chain.stream(user_q):
                    final.append(chunk)
                    holder.markdown("".join(final))
                answer = "".join(final)
            st.session_state.messages.append(("assistant", answer))